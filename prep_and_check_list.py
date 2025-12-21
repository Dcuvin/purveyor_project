from datetime import date
import pandas as pd
import sqlite3
from openpyxl import load_workbook #imports python library for reading and writting excel files
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
import sys #import sys modulet o access command-line arguments
import os #This statement is used to include the functionality of
#the os module, allowing you to interact with the operating system in a portable way
from docx import Document
from excel_format import format_headers_and_borders, set_print_options, insert_blank_rows, format_order_sheet, format_table
from prep_req import req_prep, req_prep_ver_2
from collections import defaultdict
from database import create_df

#----------------------------------------------------------------------------
def excel_prep_list(item_id, event_name, guest_count, event_start, event_date, event_location, db):
    
    current_date = date.today()
    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")
    #It will then query a junction table and pull all procedures associated with the id.          
    mise_list = []
    unique_item_names = []
    mise_list_final= []
    for id in item_id:
        cursor.execute(f"""
                       SELECT menu_prep_list.item_name, prep_list.prep
                       FROM prep_list
                       JOIN menu_prep_list ON prep_list.prep_id = menu_prep_list.prep_id
                       WHERE menu_prep_list.menu_item_id = {id};
                       """)  
    

        
        #.fetchall() is a list of tuples
        mise = cursor.fetchall()
        # access the tuple inside the list
        for tuple_item in mise:
            
            mise_list.append({'Item':str(tuple_item[0]), 'Mise':tuple_item[1], 'Need':'  '})
            # Create a list of items
            if tuple_item[0] not in unique_item_names:
                unique_item_names.append(tuple_item[0])
    conn.close()

    
    # Create a dict of items with a list of mise
    for name in unique_item_names: 
        mise_list_final.append({'Item': name, 'Mise':[], 'Need':' '})

    # Iteratively add the mise from mise_list to mise_list_2
    for item_1 in mise_list:
        for item_2 in mise_list_final:
            if item_1['Item'] == item_2['Item']:
                item_2['Mise'].append(item_1['Mise'].capitalize())

    # Iteratively title() each item name

    for item in mise_list_final:
        item['Item'] = item['Item'].title()

    print(mise_list_final)

    # Function that creates a dataframe
    def create_df(data):

        df= pd.DataFrame(data)
        return df

    df_list =[] 
    for dict_item in mise_list_final:
        df_list.append(create_df(dict_item))

    print(df_list)

    pivot_list = []
    def create_pivot(data):    
        pivot = data.pivot(index='Mise', columns='Item', values='Need')

        return pivot
    for data_frame in df_list:
        pivot_list.append(create_pivot(data_frame))

    #print(pivot_list)

    excel_file_count = 0
    # Create an excel file
    excel_file = f"prep_and_checklists/{event_name}/PREPLIST_{event_name}_{formatted_date}_{excel_file_count}.xlsx"
    # Continously checks until it finds a non-existent file name
    while os.path.exists(excel_file):
        excel_file_count += 1
        # This updates the file_count, allowing for it to be checked again in the while loop
        excel_file = f"prep_and_checklists/{event_name}/PREPLIST_{event_name}_{formatted_date}_{excel_file_count}.xlsx"
    
    #print(excel_file)
    

    # Fills-out the excel file
    with pd.ExcelWriter(excel_file, engine='openpyxl',mode = 'w') as writer:
        current_row = 3
        for pivot in pivot_list:
            pivot.to_excel(writer, sheet_name= 'prep_sheet', startrow=current_row, startcol=0)
            current_row += len(pivot) + 2 # Add space between tables

    # Load the workbook and access the sheet
    workbook = load_workbook(excel_file)
    prep_sheet= workbook["prep_sheet"]

    # Format the tables in the file
    start_row = 4
    start_col = 1
    for df in pivot_list:
        insert_blank_rows(prep_sheet, start_row)
        start_row += 1
        format_headers_and_borders(prep_sheet, start_row, start_col, 2)
        start_row += len(df) + 2

    # Insert Event Info
    title = prep_sheet.cell(row=1, column=1, value=f"{event_name} {guest_count} Guests {event_start} {event_date}")
    title.font = Font(name='Calibri', size=16, bold=True, underline='single', color='000000')
    event_info = prep_sheet.cell(row=2, column=1, value =f"Location: {event_location}")
    event_info.font = Font(name='Calibri', size=16, bold=True, underline='single', color='000000')

    # Set print options
    set_print_options(prep_sheet)

    # Save the workbook with formatting
    workbook.save(excel_file)

    # Load the workbook and select the active worksheet
    workbook = load_workbook(excel_file)
    prep_sheet= workbook["prep_sheet"]
    #---------------------------------------------------------------------------------

    # Iterate over each row and column in the sheet
    for row in prep_sheet.iter_rows():
        for cell in row:
            # Check if the cell contains 'Mise'
            if cell.value and isinstance(cell.value, str) and 'Mise' in cell.value:
                right_cell = prep_sheet.cell(row=cell.row, column=cell.column + 1)
                if right_cell.value:
                    # Replace the cell with 'Mise' with the content of the immediate right cell
                    cell.value = right_cell.value
                    # Replace the content of the right cell with 'Need'
                    right_cell.value = 'Need'

    
    
    # Save the workbook with formatting
    workbook.save(excel_file)
    
    # Fill out order_sheet
    get_order_list(item_id,db,excel_file,event_name,guest_count,event_date)

    
    print("✅ Excel Prep and Order List Created and Reformatted!")
    #return excel_file

#---------------------------------------------------------------------------------
def word_prep_list(item_id, event_name, guest_count, event_start, event_date, event_location,db):
    
    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")

    #the updated version will take a list of menu_item_ids
    #It will then query a junction table and pull all procedures associated with the id.          
    procedure_list = []
    unique_item_list =[]
    final_proc_list =[]
    for i in item_id:
        cursor.execute(f""" 
                        SELECT menu_items.item_name, procedures.item_procedure
                        FROM menu_procedures
                        JOIN menu_items ON menu_procedures.menu_item_id = menu_items.menu_item_id
                        JOIN procedures ON procedures.proc_id = menu_procedures.proc_id
                        WHERE menu_procedures.menu_item_id = {i};
            
                        """)   
    
        #.fetchall() is a list of tuples
        procedures = cursor.fetchall()
        # access the tuple inside the list
        for tuple_item in procedures:
            procedure_list.append({'item':tuple_item[0], 'proc': tuple_item[1]})
            if tuple_item[0] not in unique_item_list:
                unique_item_list.append(tuple_item[0])

    for item in unique_item_list:
        final_proc_list.append({'item':item, 'proc':[]})

    for proc_1 in procedure_list:
        for proc_2 in final_proc_list:
            if proc_1['item'] == proc_2['item']:
                proc_2['proc'].append(proc_1['proc']) 

    #print(final_proc_list)
    
                
    # Create a new Word document
    file_count = 0
    doc = Document()
    doc.add_heading(f'{event_name} {event_date}', 0)
    doc.add_heading(f'Guests: {guest_count}', level=2)
    doc.add_heading(f'Start: {event_start}', level=2)

    # Create datetime variable
    current_date = date.today()
    
    for dictionary in final_proc_list:
        doc.add_heading(f"{dictionary['item']}", level=2)

        for proc in dictionary['proc']:
            doc.add_paragraph(proc.capitalize() +' ' + '\u2610', style='List Bullet')
        
    
    # Check for any duplicate html files
    docx_file_count = 0

    prep_list_file_path = f'prep_and_checklists/{event_name}/{event_name}_Prep List_{current_date}_{docx_file_count}.docx'
    
    # continously checks until it finds a non-existent file name
    while os.path.exists(prep_list_file_path):
        docx_file_count += 1
        # this updates the file_count, allowing for it to be checked again in the while loop
        prep_list_file_path = f'prep_and_checklists/{event_name}/{event_name}_Prep List_{current_date}_{docx_file_count}.docx'

    
    doc.save(prep_list_file_path)
    print("Prep List Created!")
    

    conn.close()
#------------------------------------------------------------------------------------------
def word_checklist(item_id, event_name, guest_count, event_start, event_date, event_location,db , station_ids):
    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")

    #the updated version will take a list of menu_item_ids
    #It will then query a junction table and pull all procedures associated with the id.          
    mise_list = []
    unique_item_list =[]
    final_mise_list =[]
    for i in item_id:
        cursor.execute(f""" 
                        SELECT menu_mise_checklist.item_name, mise_checklist.mise_en_place
                        FROM menu_mise_checklist
                        JOIN mise_checklist ON menu_mise_checklist.checklist_id = mise_checklist.checklist_id
                        WHERE menu_mise_checklist.menu_item_id = {i};
            
                        """)   
    
        #.fetchall() is a list of tuples
        mise_en_place = cursor.fetchall()
        # access the tuple inside the list
        for tuple_item in mise_en_place:
            mise_list.append({'item':tuple_item[0], 'mise': tuple_item[1]})
            if tuple_item[0] not in unique_item_list:
                unique_item_list.append(tuple_item[0])

    for item in unique_item_list:
        final_mise_list.append({'item':item, 'mise':[]})

    for mise_1 in mise_list:
        for mise_2 in final_mise_list:
            if mise_1['item'] == mise_2['item']:
                mise_2['mise'].append(mise_1['mise']) 

    # capitalize the first letter 
    for menu_item in final_mise_list:
        menu_item['item'] = menu_item['item'].title()
        for mise in menu_item['mise']:
            mise.title()

    #print(final_proc_list)
    final_mise_list.append({'item': 'Dry Goods/Tools', 'mise':['Maldon','EVOO','C-folds','Vodka Spray','Quarter Sheet Trays','Half Sheet Trays','Catering Trays', 'Cutting boards', 'Mixing Bowls', 'Sani-wipes','Gloves', 'Tasting Spoons','Piping Bags', 'Quarts','Pints', 'Lids']})
    # Adding dry-goods/ tools section to checklist
           
    # Create a new Word document
    file_count = 0
    doc = Document()
    doc.add_heading(f'{event_name} {event_date}', 0)
    doc.add_heading(f'Guests: {guest_count}', level=2)
    doc.add_heading(f'Start: {event_start}', level=2)
    doc.add_heading(f'Loction: {event_location}', level = 2)

    # Create datetime variable
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")

    for dict in final_mise_list:
        doc.add_heading(f"{dict['item']}", level=2)

        # Add items as paragraphs with a checkbox
        for mise in dict['mise']:
            doc.add_paragraph('\u2610' + ' ' + mise.capitalize())
            
        
    
    # Check for any duplicate files
    docx_file_count = 0

    checklist_file_path = f'prep_and_checklists/{event_name}/CHECKLIST_{event_name}_{formatted_date}_{docx_file_count}.docx'
    
    #continously checks until it finds a non-existent file name
    while os.path.exists(checklist_file_path):
        file_count += 1
        # this updates the file_count, allowing for it to be checked again in the while loop
        checklist_file_path = f'prep_and_checklists/{event_name}/CHECKLIST_{event_name}_{formatted_date}_{file_count}.docx'

    
    doc.save(checklist_file_path)
    print("✅  Checklist Created!")
    

    conn.close()
#------------------------------------------------------------------------------------------
def get_order_list(item_id,db,excel_file_path,event_name,guest_count,event_date):

    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")

    result_dict= {'Ingredient': [],'QTY':'', 'Purveyor':[]}
    for id in item_id:
        cursor.execute("""
                SELECT ingredients.ingredient_name, ingredients.purveyor
                FROM ingredients
                JOIN menu_ingredients ON ingredients.ingredient_id = menu_ingredients.ingredient_id
                WHERE menu_ingredients.menu_item_id = ?;
            """, (id,)
        )
        
        results = cursor.fetchall()
        #print(results)

        # Remove duplicate ingredients
        for tuple_item in results:
            capitalized_ingredient = tuple_item[0].capitalize()
            if tuple_item[0].capitalize() not in result_dict['Ingredient']:
                result_dict['Ingredient'].append(capitalized_ingredient)
                result_dict['Purveyor'].append(tuple_item[1].capitalize())

    #print(result_list)

    print(result_dict)
    # Function that creates a dataframe
    def create_df(data):

        df= pd.DataFrame(data)
        return df
    
    
    df_list =[] 
    df_list.append(create_df(result_dict))

    print(df_list)

   
    # Create the order_sheet and populate with data.

    with pd.ExcelWriter(excel_file_path, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
        sheet_name = "order_sheet"
        pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)

        # Add event info to the top of the order sheet.

        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        worksheet["A1"] = f"Order List for {event_name} - {event_date} - Guest:{guest_count}"  # Customize your title here

        current_row = 2
        for item in df_list:
            item.to_excel(writer, sheet_name= 'order_sheet', startrow=current_row, startcol=0, index=False)
            current_row += len(item.index) + 1  # Increment to avoid overlap

    # Reload the workbook with openpyxl to apply the formatting.
    workbook = load_workbook(excel_file_path)


    format_order_sheet(workbook['order_sheet'], 2, 1, 3)

    workbook.save(excel_file_path)
#----------------------------------------------------------------------------


    print("Order Sheet Created!")

def get_order_list_ver_2(item_id,db,excel_file_path,event_name,guest_count,event_date):

    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")

    result_dict= {'Ingredient': [],'QTY':'', 'Purveyor':[]}
    for id in item_id:
        cursor.execute("""
                SELECT ingredients.ingredient_name, ingredients.purveyor
                FROM ingredients
                JOIN menu_ingredients ON ingredients.ingredient_id = menu_ingredients.ingredient_id
                WHERE menu_ingredients.menu_item_id = ?;
            """, (id,)
        )
        
        results = cursor.fetchall()
        #print(results)

        # Remove duplicate ingredients
        for tuple_item in results:
            capitalized_ingredient = tuple_item[0].capitalize()
            if tuple_item[0].capitalize() not in result_dict['Ingredient']:
                result_dict['Ingredient'].append(capitalized_ingredient)
                result_dict['Purveyor'].append(tuple_item[1].capitalize())

    #print(result_list)

    print(result_dict)
    # Function that creates a dataframe
    def create_df(data):

        df= pd.DataFrame(data)
        return df
    
    
    df_list =[] 
    df_list.append(create_df(result_dict))

    print(df_list)

   
    # Create the order_sheet and populate with data.

    with pd.ExcelWriter(excel_file_path, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
        sheet_name = "order_sheet"
        pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)

        # Add event info to the top of the order sheet.

        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        worksheet["A1"] = f"Order List for {event_name} - {event_date} - Guest:{guest_count}"  # Customize your title here

        current_row = 2
        for item in df_list:
            item.to_excel(writer, sheet_name= 'order_sheet', startrow=current_row, startcol=0, index=False)
            current_row += len(item.index) + 1  # Increment to avoid overlap

    # Reload the workbook with openpyxl to apply the formatting.
    workbook = load_workbook(excel_file_path)


    format_order_sheet(workbook['order_sheet'], 2, 1, 3)

    workbook.save(excel_file_path)


    print("Order Sheet Created!")
    
    
#----------------------------------------------------------------------------

def old_excel_prep_list(item_id, event_name, guest_count, event_start, event_date,db):
    pass

#----------------------------------------------------------------------------
def old_word_checklist(item_id, event_name, guest_count, event_start, event_date,db):
    pass

#----------------------------------------------------------------------------

def excel_prep_list_ver_2(item_id, event_name, guest_count, event_start, event_date, event_location, db, station_ids, event_type):


    current_date = date.today()
    conn = sqlite3.connect(db)
    # Cursor to execute commands
    cursor = conn.cursor()
    current_date = date.today()
    formatted_date = current_date.strftime("%m-%d-%Y")
    #It will then query a junction table and pull all procedures associated with the id. 
    item_ids_list = item_id      
    menu_item_list =[]
    #station_menu_items_list =[]   
    stations = []
    station_menu_ids = []
    
    
    if station_ids:
        for id in station_ids:
            cursor.execute(f"""
                        SELECT menu_items_stations.station_name, menu_items_stations.menu_item_id 
                        FROM menu_items_stations
                        WHERE menu_items_stations.station_id = {id};
                        """)
            #.fetchall() is a list of tuples
            station_menu_items = cursor.fetchall()
            #print(f"station_menu_items: {station_menu_items}")
            # access the tuple inside the list
            for tuple_item in station_menu_items:
                stations.append({"station_name": tuple_item[0],"menu_item_id": tuple_item[1],"menu_item_name":"","mise": []})
        
                station_menu_ids.append(tuple_item[1])
        #print(f"station_menu_ids: {station_menu_ids}")  
        for station_dict in stations:
            cursor.execute(f"""
                        SELECT menu_prep_list.menu_item_id, menu_prep_list.item_name, prep_list.prep
                        FROM prep_list
                        JOIN menu_prep_list ON prep_list.prep_id = menu_prep_list.prep_id
                        WHERE menu_prep_list.menu_item_id = {station_dict['menu_item_id']};
                        """)
            station_mise = cursor.fetchall()
            print(station_mise)
            for mise_tuple in station_mise:
                if mise_tuple[0] == station_dict['menu_item_id']:
                    station_dict['menu_item_name'] = mise_tuple[1]
                    station_dict['mise'].append(mise_tuple[2])

        
        for station_dict in stations:
            menu_item_list.append({'Item':station_dict["menu_item_name"], 'Category':station_dict["station_name"],'Mise':station_dict["mise"], 'Need':'  '})
    
    for id in item_id:
        
        cursor.execute(f"""
                       SELECT menu_items.item_name, menu_items.category
                       FROM menu_items
                       WHERE menu_items.menu_item_id = {id};
                       """)  
         #.fetchall() is a list of tuples
        menu_items = cursor.fetchall()
        # access the tuple inside the list
        for tuple_item in menu_items:
            
            menu_item_list.append({'Item':str(tuple_item[0]), 'Category':tuple_item[1],'Mise':[], 'Need':'  '})
    
    
    for id in item_id:
        cursor.execute(f"""
                       SELECT menu_prep_list.item_name, prep_list.prep
                       FROM prep_list
                       JOIN menu_prep_list ON prep_list.prep_id = menu_prep_list.prep_id
                       WHERE menu_prep_list.menu_item_id = {id};
                       """)  

        
        #.fetchall() is a list of tuples
        mise = cursor.fetchall()
        print(f"mise: {mise}")
        # access the tuple inside the list
        for tuple_item in mise:
            for menu_item in menu_item_list:
                if str(tuple_item[0]) == menu_item['Item']:
                        if str(tuple_item[1]) not in menu_item['Mise']:
                            menu_item['Mise'].append(tuple_item[1])
            
    conn.close()

    # Iteratively title() each item name

    for item in menu_item_list:
        item['Item'] = item['Item'].title()
        for i in range(len(item['Mise'])):
            item['Mise'][i].title()
            
    #print(f"menu_item_list: {menu_item_list}")

    


    # Count how many categories there are and create that many tables in the excel file.
    category_list = []
    for item in menu_item_list:
        if item['Category'] not in category_list:
            category_list.append(item['Category'])
        else:
            continue
    
    category_count = len(category_list)
            
    #print(f"Number of Categories: {category_count}")

    # Create a list df based on menu_item category

    df_list = [{'Category': category, 'DataFrame': []} for category in category_list]
    for menu_item in menu_item_list:
        for df_item in df_list:
            if menu_item['Category'] ==  df_item['Category']:
                df_item['DataFrame'].append(create_df(menu_item).pivot(index='Mise', columns='Item', values='Need'))
    
    #title each category
    for dict_item in df_list:
        dict_item["Category"].title()

    print(f"df_list: {df_list}")
    
    excel_file_count = 0
    # Create an excel file
    excel_file = f"prep_and_checklists/{event_name}/PREPLIST_{event_name}_{formatted_date}_{excel_file_count}.xlsx"
    # Continously checks until it finds a non-existent file name
    while os.path.exists(excel_file):
        excel_file_count += 1
        # This updates the file_count, allowing for it to be checked again in the while loop
        excel_file = f"prep_and_checklists/{event_name}/PREPLIST_{event_name}_{formatted_date}_{excel_file_count}.xlsx"
        
    #print(excel_file)
        
    # Fills-out the excel file
    with pd.ExcelWriter(excel_file, engine='openpyxl',mode = 'w') as writer:
        category_row = 3
        timing_row = 4
        df_row = 4
        current_col = 0  # 0-based for pandas
        current_category_col = 1 #1-based for openpyxl (col A)
        current_category_end_col = 2 #2-based for openpyxl (col B)
        
        # Create the sheet by writing a placeholder DataFrame, as writer.sheets doesn't work if no sheet exists.
        pd.DataFrame().to_excel(writer, sheet_name='prep_sheet', index=False)
        for df_dict in df_list:
         
            current_category_col_letter = get_column_letter(current_category_col)
            end_category_col_letter = get_column_letter(current_category_end_col)

        
            ws = writer.sheets['prep_sheet']

            # merge cells for category title
            ws.merge_cells(f"{current_category_col_letter}{category_row}:{end_category_col_letter}{category_row}")
            

            # Insert category as title 
            cell_category = ws[f"{current_category_col_letter}{category_row}"]
            cell_category.value = df_dict['Category'].capitalize()
            cell_category.font = Font(bold=True, name='Calibri', size=14, color="000000")
            cell_category.fill = PatternFill(start_color="FFC9DAF8", end_color="FFC9DAF8", fill_type="solid")
            cell_category.border = Border(left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin'))
            cell_category.alignment = Alignment(horizontal='center', vertical='center')



            # merge cells for timing info
            ws.merge_cells(f"{current_category_col_letter}{timing_row}:{end_category_col_letter}{timing_row}")

            # Insert, format, and merge row below the 'category' title for timing
            cell_timing= ws[f"{current_category_col_letter}{timing_row}"]
            cell_timing.value = "Timing TBD"
            cell_timing.font = Font(bold=True, name='Calibri', size=14, color="FFFF0000")
            cell_timing.border = Border(left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin'))
            cell_timing.alignment = Alignment(horizontal='center', vertical='center')


            for dataframe in df_dict['DataFrame']:
                #this is pandas based and starts at 0 for indexing
                dataframe.to_excel(writer, sheet_name= 'prep_sheet', startrow=df_row, startcol=current_col)
                format_table(ws, df_row, current_category_col, dataframe)
                #this is openpyxl based and starts at 1 for indexing
                df_row += len(dataframe) + 2 # Add space between tables
                #insert_blank_rows(ws, df_row + 1 )
            df_row = 4
            current_col += 3
            current_category_col += 3
            current_category_end_col += 3
      

        

     # Load the workbook and access the sheet
    workbook = load_workbook(excel_file)
    prep_sheet= workbook["prep_sheet"]

    # Save the workbook with formatting
    workbook.save(excel_file)
    #(print('Workbook Created!'))



    # Insert Event Info
    title = prep_sheet.cell(row=1, column=1, value=f"{event_name}, Guests: {guest_count} , {event_start} ,{event_date}")
    title.font = Font(name='Calibri', size=16, bold=True, underline='single', color='000000')
    event_info = prep_sheet.cell(row=2, column=1, value =f"Location: {event_location}")
    event_info.font = Font(name='Calibri', size=16, bold=True, underline='single', color='000000')

    # Set print options
    set_print_options(prep_sheet)

    # Save the workbook with formatting
    workbook.save(excel_file)

    # Load the workbook and select the active worksheet
    workbook = load_workbook(excel_file)
    prep_sheet= workbook["prep_sheet"]
    #---------------------------------------------------------------------------------

    # Iterate over each row and column in the sheet
    for row in prep_sheet.iter_rows():
        for cell in row:
            # Check if the cell contains 'Mise'
            if cell.value and isinstance(cell.value, str) and 'Mise' in cell.value:
                right_cell = prep_sheet.cell(row=cell.row, column=cell.column + 1)
                if right_cell.value:
                    # Replace the cell with 'Mise' with the content of the immediate right cell
                    cell.value = right_cell.value
                    # Replace the content of the right cell with 'Need'
                    right_cell.value = 'Need'

    
    
    # Save the workbook with formatting
    workbook.save(excel_file)
    
    # Consolidate item_ids from both individual menu_items and each item within stations

    final_ids = []

    for id in station_menu_ids:
        final_ids.append(id)
    for id in item_ids_list:
        final_ids.append(id)
    print(f"final_item_ids:{final_ids}")
    # Fill out order_sheet
    get_order_list(final_ids,db,excel_file,event_name,guest_count,event_date)

    #req_prep_ver_2(final_ids, new_folder_path, event_date, event_name,db)

    print("✅ Excel Prep and Order List Created and Reformatted!")
    return final_ids
